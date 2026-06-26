from copy import deepcopy
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile
import re

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parent
TEMPLATE = ROOT / "PM02_PP_template.docx"
OUT = ROOT / "Отчетная_документация_по_ПП_ПМ02_BASIC_исправлено.docx"

FULL_NAME = "Драпов Матвей Борисович"
FULL_NAME_GEN = "Драпова Матвея Борисовича"
GROUP_DOTS = "01-23.ИСИП.ОФ.9"
GROUP_SPACES = "01-23 ИСИП ОФ 9"
COURSE = "3"
SPECIALITY = "09.02.07 «Информационные системы и программирование»"
PLACE = "АНПОО «Хекслет колледж»"
ADDRESS_TITLE = "190020, г. Санкт-Петербург, Рижский проспект, дом № 8, литер А"
PERIOD = "25.05.2026 - 14.06.2026"
ORG_SUPERVISOR = "Васильев Антон Александрович"


REPORT_BLOCKS = {
    "ОПИСАТЬ АРХИТЕКТУРУ": [
        "В рамках BASIC-1 была изучена документация archtool и web_fractal. Были разобраны DependencyInjector, AppModule, Layer, двухпроходный алгоритм инъекции зависимостей, UnitOfWork, BaseRepo и HttpControllerABC. По итогам был подготовлен markdown-документ с ответами своими словами и примерной структурой проекта.",
        "В BASIC-2 был разработан мини-проект ToDo API. Архитектура построена по двум bounded context: users отвечает за регистрацию и получение пользователей, todos - за CRUD задач, привязанных к пользователю. В каждом модуле сохранена структура interfaces.py, models.py, repos.py, services.py и controllers.py.",
        "Взаимодействие слоев организовано через интерфейсы. Репозитории реализуют UserRepoABC и TodoRepoABC, сервисы реализуют UserServiceABC и TodoServiceABC, контроллеры обращаются к сервисам через абстракции. TodoService дополнительно зависит от UserServiceABC, чтобы проверять существование пользователя без прямой зависимости от конкретной реализации модуля users.",
        "Общая схема работы приложения: entrypoints/run.py создает FastAPI-приложение, app/archtool_conf/bundle_project.py собирает DependencyInjector, engine и session_maker, после чего archtool находит реализации интерфейсов, а web_fractal подключает контроллеры к FastAPI. Запрос проходит по цепочке controller - service - repo - SQLite.",
    ],
    "ОПИСАТЬ МЕХАНИЗМЫ ИНТЕГРАЦИИ": [
        "В BASIC-2 была выполнена интеграция archtool, web_fractal, FastAPI, SQLAlchemy async, aiosqlite и uvicorn. FastAPI использовался для HTTP API, SQLAlchemy - для ORM и работы с базой данных, web_fractal - для UnitOfWork, репозиториев и подключения HTTP-контроллеров.",
        "DependencyInjector был настроен с явной передачей слоев приложения и AppModule для app.users и app.todos. session_maker зарегистрирован через injector.register(), чтобы репозитории получали фабрику сессий через DI, а не через ручную передачу в конструкторах.",
        "Интеграция с web_fractal выполнялась через import_all_models() и initialize_controllers_api(). Первая функция нужна для загрузки ORM-моделей перед созданием таблиц, вторая подключает контроллеры к FastAPI. Контроллеры наследуются от HttpControllerABC и регистрируют маршруты в init_http_routes().",
        "Для запуска и проверки были подготовлены run.ps1 и check.ps1. Первый скрипт запускает приложение для демонстрации Swagger UI, второй автоматически поднимает сервер, выполняет curl-запросы ко всем основным эндпоинтам и останавливает процесс.",
    ],
    "ОПИСАТЬ ПРОЦЕСС ОТЛАДКИ": [
        "При выполнении BASIC-2 и BASIC-3 проводилась отладка запуска приложения, регистрации зависимостей и работы HTTP-эндпоинтов. Использовались PowerShell, uv, uvicorn, Swagger UI, curl, server.out.log, server.err.log и проверка SQLite-файла todo_app.db.",
        "В ходе отладки были выявлены особенности совместной работы актуальных версий archtool и web_fractal. Для стабильного запуска были добавлены недостающие зависимости и проверены места, связанные с импортом моделей и подключением контроллеров. После исправлений приложение стало запускаться одной командой.",
        "Проверялась работа UnitOfWork в репозиториях: создание пользователя, получение пользователя, создание задачи, получение списка задач и отметка задачи выполненной. Данные сохранялись в SQLite-базе, что подтверждало корректную работу транзакций и репозиториев.",
        "В BASIC-5 дополнительно была проверена среда разработки: WSL/Ubuntu, uv/Python, FastAPI, Node.js/npm, Docker и VS Code. Это позволило подготовить рабочее окружение для дальнейших задач.",
    ],
    "ОПИСАТЬ, КАКИЕ ТЕСТОВЫЕ НАБОРЫ": [
        "Для BASIC-2 и BASIC-3 были составлены ручные тестовые сценарии для всех обязательных эндпоинтов ToDo API. В качестве тестовых данных использовались пользователь demo@example.com с именем Demo User и задачи, привязанные к созданному user_id.",
        "Основной сценарий включал POST /users/, GET /users/{id}, POST /todos/, GET /todos/?user_id=... и PATCH /todos/{id}/complete. Так проверялись создание пользователя, сохранение данных, создание задачи, получение списка задач и изменение статуса completed.",
        "Дополнительно проверялась целостность связи между users и todos: задача создавалась только для существующего пользователя, а список задач фильтровался по user_id. Проект также запускался на чистом виртуальном окружении через uv, чтобы убедиться в воспроизводимости README и скриптов.",
        "В BASIC-3 был подготовлен документ для ревью и демо: план запуска, порядок демонстрации Swagger UI и ответы на ожидаемые вопросы по DI, слоям, UnitOfWork и добавлению нового модуля.",
    ],
    "ОПИСАТЬ, КАКИЕ ШАБЛОНЫ": [
        "В проекте использовались принципы, изученные в BASIC-1 и BASIC-4: разделение ответственности, dependency inversion, interface-first design и слоистая архитектура. Контроллеры отвечают за HTTP-запросы, сервисы - за бизнес-логику, репозитории - за доступ к данным, интерфейсы - за контракты между слоями.",
        "Для доступа к данным применялся шаблон Repository: UserRepo и TodoRepo скрывают детали SQLAlchemy от сервисов. Для транзакций использовался Unit of Work: UnitOfWork открывает сессию, выполняет commit при успешной операции и rollback при ошибке.",
        "Код оформлен по практикам Python-проектов: классы названы в PascalCase, функции и переменные - в snake_case, интерфейсы вынесены в interfaces.py, ORM-модели - в models.py, точка входа - в entrypoints/run.py. Зависимости объявлены через аннотации классов, как требует archtool.",
        "В BASIC-3 было проведено повторное чтение ключевых файлов проекта и подготовлены ответы на вопросы ревью. Это подтвердило понимание архитектуры, а не только механическое копирование структуры.",
        "Отчетная документация оформлена по требованиям инструкции: Times New Roman 12, полуторный интервал, выравнивание по ширине, поля 30 мм слева, 10 мм справа и 20 мм сверху и снизу. Цветовые выделения в заполненных областях удалены.",
    ],
}


def set_runs_text(paragraph, text):
    if not paragraph.runs:
        run = paragraph.add_run(text)
    else:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
        run = paragraph.runs[0]
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)


def set_cell(cell, text):
    set_runs_text(cell.paragraphs[0], text)


def paragraph_after(paragraph, text):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if paragraph._p.pPr is not None:
        new_para._p.insert(0, deepcopy(paragraph._p.pPr))
    set_runs_text(new_para, text)
    return new_para


def delete_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def strip_color_only(path: Path):
    tmp = path.with_suffix(".tmp.docx")
    with ZipFile(path, "r") as zin, ZipFile(tmp, "w", ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename.startswith("word/") and item.filename.endswith(".xml"):
                text = data.decode("utf-8")
                text = re.sub(r"<w:highlight\b[^>]*/>", "", text)
                text = re.sub(r"<w:shd\b[^>]*/>", "", text)
                data = text.encode("utf-8")
            zout.writestr(item, data)
    tmp.replace(path)


def fill():
    doc = Document(str(TEMPLATE))

    paragraph_replacements = {
        "Место прохождения практики:__________________________________________________": f"Место прохождения практики: {PLACE}",
        "____________________________________________________________________________": ADDRESS_TITLE,
        "_______________________________ (___________________)": f"{ORG_SUPERVISOR} (___________________)",
        "(___________________) « 14 » мая 2026г.": f"{ORG_SUPERVISOR} (___________________) « 14 » июня 2026 г.",
        "« 14 » мая 2026г.": "« 14 » июня 2026 г.",
        "Заключение: обучающийся _______________________________________________ готов(а) к выполнению профессиональной деятельности по Осуществлению интеграции программных модулей.": f"Заключение: обучающийся {FULL_NAME} готов к выполнению профессиональной деятельности по осуществлению интеграции программных модулей.",
        "Обучающийся _______________________________(________________)": f"Обучающийся {FULL_NAME} (________________)",
    }
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text in paragraph_replacements:
            set_runs_text(paragraph, paragraph_replacements[text])

    t0 = doc.tables[0]
    for row, value in [(0, SPECIALITY), (4, FULL_NAME), (5, COURSE), (7, GROUP_DOTS), (8, PERIOD)]:
        set_cell(t0.cell(row, 1), value)

    for idx in [2, 5, 7]:
        table = doc.tables[idx]
        for row, value in [
            (0, FULL_NAME_GEN),
            (1, SPECIALITY),
            (2, COURSE),
            (3, GROUP_SPACES),
            (5, PLACE),
            (6, PERIOD),
        ]:
            set_cell(table.cell(row, 1), value)

    t3 = doc.tables[3]
    for row, value in [
        (0, FULL_NAME),
        (1, COURSE),
        (2, SPECIALITY),
        (3, GROUP_SPACES),
        (4, "Производственная"),
        (5, PLACE),
        (6, PERIOD),
    ]:
        set_cell(t3.cell(row, 1), value)

    for row_idx in range(1, 16):
        set_cell(doc.tables[4].cell(row_idx, 0), str(row_idx))

    for row_idx in range(1, 12):
        set_cell(doc.tables[8].cell(row_idx, 2), "+")

    t9 = doc.tables[9]
    for row, value in [
        (0, FULL_NAME),
        (1, SPECIALITY),
        (2, COURSE),
        (3, "Осуществление интеграции программных модулей"),
        (4, GROUP_SPACES),
        (5, PLACE),
        (6, PERIOD),
    ]:
        set_cell(t9.cell(row, 1), value)

    for paragraph in list(doc.paragraphs):
        raw = paragraph.text.strip()
        key = next((k for k in REPORT_BLOCKS if raw.startswith(k)), None)
        if key:
            inserted = paragraph
            set_runs_text(inserted, REPORT_BLOCKS[key][0])
            for item in REPORT_BLOCKS[key][1:]:
                inserted = paragraph_after(inserted, item)

    doc.core_properties.author = FULL_NAME
    doc.core_properties.last_modified_by = FULL_NAME
    doc.save(str(OUT))
    strip_color_only(OUT)
    print(OUT)


if __name__ == "__main__":
    fill()
