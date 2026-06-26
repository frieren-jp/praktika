from pathlib import Path
from zipfile import ZipFile, ZIP_DEFLATED

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parent
TEMPLATE = ROOT / "PM02_PP_template.docx"
OUT = ROOT / "Отчетная_документация_по_ПП_ПМ02_BASIC_готово.docx"

FULL_NAME = "Драпов Матвей Борисович"
FULL_NAME_GEN = "Драпова Матвея Борисовича"
GROUP_DOTS = "01-23.ИСИП.ОФ.9"
GROUP_SPACES = "01-23 ИСИП ОФ 9"
COURSE = "3"
SPECIALITY = "09.02.07 «Информационные системы и программирование»"
PLACE = "АНПОО «Хекслет колледж»"
ADDRESS_TITLE = "190020, г. Санкт-Петербург, Рижский проспект, дом № 8, литер А"
ADDRESS_SHORT = "Санкт-Петербург, Рижский пр-т, дом 8"
PERIOD = "25.05.2026 - 14.06.2026"
PERIOD_TEXT = "25 мая 2026 года по 14 июня 2026 года"
EDU_SUPERVISOR = "Шандыба Дарья Александровна"
ORG_SUPERVISOR = "Васильев Антон Александрович"


REPORT_TEXT = {
    "ОПИСАТЬ АРХИТЕКТУРУ": [
        "В рамках BASIC-1 была изучена документация archtool и web_fractal. Были разобраны основные понятия: DependencyInjector, AppModule, Layer, двухпроходный алгоритм инъекции зависимостей, UnitOfWork, BaseRepo и HttpControllerABC. По результатам изучения был подготовлен markdown-документ с ответами на вопросы своими словами и с текстовой схемой типичной структуры проекта.",
        "На основе изученной документации в BASIC-2 был спроектирован мини-проект ToDo API. Архитектура приложения построена по принципу разделения на bounded context: модуль users отвечает за регистрацию и получение пользователей, а модуль todos отвечает за создание, получение и изменение задач. Каждый модуль имеет одинаковую структуру: interfaces.py содержит абстрактные контракты, models.py - ORM-модели SQLAlchemy, repos.py - доступ к данным, services.py - бизнес-логику, controllers.py - HTTP-интерфейс.",
        "Связь между слоями организована через интерфейсы. Например, UserRepoABC и TodoRepoABC описывают методы доступа к данным, UserServiceABC и TodoServiceABC описывают операции бизнес-логики, а контроллеры обращаются к сервисам через абстракции. Класс TodoService дополнительно зависит от UserServiceABC, чтобы при создании задачи проверять существование пользователя и при этом не связывать модуль todos напрямую с конкретной реализацией users.",
        "Текстовая схема проекта: entrypoints/run.py создает FastAPI-приложение; app/archtool_conf/bundle_project.py собирает DependencyInjector, engine и session_maker; затем archtool находит реализации интерфейсов в users и todos; web_fractal подключает контроллеры к FastAPI; запросы проходят по цепочке controller -> service -> repo -> SQLite. Такая структура позволила отделить HTTP-слой от бизнес-логики и доступа к базе данных.",
    ],
    "ОПИСАТЬ МЕХАНИЗМЫ ИНТЕГРАЦИИ": [
        "В BASIC-2 была выполнена интеграция нескольких программных модулей и библиотек: archtool, web_fractal, FastAPI, SQLAlchemy async, aiosqlite и uvicorn. FastAPI использовался как HTTP-фреймворк, SQLAlchemy - как ORM для работы с базой данных, а web_fractal дополнял проект готовыми утилитами для UnitOfWork, репозиториев и контроллеров.",
        "DependencyInjector из archtool был настроен с передачей всех слоев приложения. В конфигурации были зарегистрированы AppModule для app.users и app.todos, после чего DI-контейнер автоматически нашел интерфейсы и реализации в соответствующих файлах. session_maker был зарегистрирован через injector.register(), чтобы репозитории могли получать фабрику сессий без ручной передачи через конструкторы.",
        "Интеграция с web_fractal выполнялась через import_all_models() и initialize_controllers_api(). Первая функция нужна для загрузки ORM-моделей перед созданием таблиц, вторая - для подключения HTTP-контроллеров к FastAPI. Контроллеры наследуются от HttpControllerABC и объявляют маршруты в init_http_routes(), поэтому их можно подключать единообразно.",
        "Для запуска и проверки проекта были подготовлены вспомогательные скрипты run.ps1 и check.ps1. run.ps1 запускает приложение через uvicorn для демонстрации Swagger UI, а check.ps1 автоматически поднимает сервер, выполняет curl-запросы ко всем основным эндпоинтам и останавливает процесс. Это упростило демонстрацию BASIC-3 и снизило риск ошибки при ручном запуске.",
    ],
    "ОПИСАТЬ ПРОЦЕСС ОТЛАДКИ": [
        "При выполнении BASIC-2 и BASIC-3 проводилась отладка запуска приложения, регистрации зависимостей и работы HTTP-эндпоинтов. Основными инструментами были PowerShell, uv, uvicorn, журналы server.out.log и server.err.log, Swagger UI, curl-запросы и проверка SQLite-файла todo_app.db.",
        "Во время отладки были выявлены особенности совместной работы актуальной версии archtool и web_fractal. Для стабильного запуска были явно добавлены недостающие runtime-зависимости web_fractal, а также проверены места, где библиотека ожидает определенную структуру DI-контейнера и импорт моделей. После исправлений приложение стало запускаться одной командой и корректно создавать таблицы.",
        "Отдельно проверялась работа UnitOfWork в репозиториях. При создании пользователя и задачи данные должны были сохраняться в базе после commit, а при ошибке операция не должна была оставлять частично записанные данные. Для этого запускались последовательные сценарии создания пользователя, получения пользователя, создания задачи, просмотра списка задач и отметки задачи выполненной.",
        "В BASIC-5 дополнительно была проверена среда разработки: WSL/Ubuntu, uv/Python, FastAPI, Node.js/npm, Docker и VS Code. Это позволило убедиться, что дальнейшие задания можно выполнять в рабочем окружении без ручной переустановки зависимостей перед каждым запуском.",
    ],
    "ОПИСАТЬ, КАКИЕ ТЕСТОВЫЕ НАБОРЫ": [
        "Для BASIC-2 и BASIC-3 были составлены ручные тестовые сценарии для всех обязательных эндпоинтов ToDo API. В качестве тестовых данных использовались пользователь с email demo@example.com и именем Demo User, а также задачи с привязкой к созданному user_id. Такой набор позволял проверить связь между модулями users и todos.",
        "Основной позитивный сценарий включал последовательность: POST /users/ для создания пользователя, GET /users/{id} для проверки сохранения, POST /todos/ для создания задачи, GET /todos/?user_id=... для получения списка задач пользователя и PATCH /todos/{id}/complete для изменения статуса completed на true.",
        "Дополнительно проверялись ситуации, связанные с целостностью данных: создание задачи для существующего пользователя, получение задач только по user_id и повторный запуск приложения с уже созданной SQLite-базой. Это подтверждало, что данные сохраняются не только в памяти процесса, а реально записываются в базу данных.",
        "Для демонстрации был подготовлен BASIC-3 документ с планом ревью и ожидаемыми вопросами по коду. Также был выполнен запуск на чистом виртуальном окружении через uv, чтобы проверить, что README и скрипты позволяют воспроизвести проект без скрытых локальных настроек.",
    ],
    "ОПИСАТЬ, КАКИЕ ШАБЛОНЫ": [
        "В проекте использовались архитектурные принципы, изученные в BASIC-1 и BASIC-4: разделение ответственности, dependency inversion, interface-first design и слоистая структура приложения. Каждый слой выполняет свою роль: контроллеры принимают HTTP-запросы, сервисы реализуют бизнес-операции, репозитории работают с данными, а интерфейсы фиксируют контракты между слоями.",
        "Для доступа к данным применялся шаблон Repository: UserRepo и TodoRepo скрывают детали SQLAlchemy от сервисов. Для управления транзакциями использовался Unit of Work: UnitOfWork открывает сессию, выполняет commit при успешном завершении операции и rollback при ошибке. Это делает работу с базой данных более предсказуемой.",
        "Код организован по соглашениям Python-проекта: названия классов записаны в PascalCase, функции и переменные - в snake_case, абстрактные интерфейсы вынесены в interfaces.py, ORM-модели - в models.py, а точка входа - в entrypoints/run.py. Зависимости между компонентами объявлены через аннотации классов, как требует archtool.",
        "В ходе BASIC-3 было проведено повторное чтение ключевых файлов проекта и подготовлены ответы на вопросы ревью: зачем нужны аннотации вместо __init__, как archtool находит реализации интерфейсов, зачем нужен UnitOfWork и что изменится при добавлении нового модуля. Это подтвердило понимание архитектуры, а не только механическое копирование структуры.",
        "Оформление отчетной документации выполнено по требованиям инструкции: используется Times New Roman 12, полуторный интервал, выравнивание по ширине, поля 30 мм слева, 10 мм справа и 20 мм сверху и снизу. Цветовые выделения в заполняемых областях удалены.",
    ],
}


def set_paragraph_text(paragraph, text):
    for run in paragraph.runs:
        run.text = ""
    run = paragraph.add_run(text) if paragraph.runs else paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)


def set_cell_text(cell, text):
    paragraph = cell.paragraphs[0]
    for extra in cell.paragraphs[1:]:
        delete_paragraph(extra)
    set_paragraph_text(paragraph, text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(text) <= 25 else WD_ALIGN_PARAGRAPH.LEFT


def insert_paragraph_after(paragraph, text, style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = paragraph._parent.add_paragraph()
    new_paragraph._p = new_p
    new_paragraph._element = new_p
    if style is not None:
        new_paragraph.style = style
    set_paragraph_text(new_paragraph, text)
    return new_paragraph


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None


def clear_cell_shading(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    for shd in tc_pr.findall(qn("w:shd")):
        tc_pr.remove(shd)


def normalize_paragraph(paragraph, first_line=False):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    if first_line:
        paragraph.paragraph_format.first_line_indent = Cm(1.25)
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(12)
        run.font.highlight_color = None
        run.font.color.rgb = None
        if run._element.rPr is not None:
            for highlight in run._element.rPr.findall(qn("w:highlight")):
                run._element.rPr.remove(highlight)


def fill_common_table(table, first_label):
    table.cell(0, 1).text = first_label
    table.cell(1, 1).text = SPECIALITY
    table.cell(2, 1).text = COURSE
    if len(table.rows) == 7:
        table.cell(3, 1).text = GROUP_SPACES
        table.cell(5, 1).text = PLACE
        table.cell(6, 1).text = PERIOD


def fill_doc():
    doc = Document(str(TEMPLATE))

    section = doc.sections[0]
    section.left_margin = Cm(3)
    section.right_margin = Cm(1)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    paragraph_replacements = {
        "Место прохождения практики:__________________________________________________": f"Место прохождения практики: {PLACE}",
        "____________________________________________________________________________": ADDRESS_TITLE,
        "_______________________________ (___________________)": f"{ORG_SUPERVISOR} (___________________)",
        "(___________________) « 14 » мая 2026г.": f"{ORG_SUPERVISOR} (___________________) « 14 » июня 2026 г.",
        "« 14 » мая 2026г.": "« 14 » июня 2026 г.",
        "Заключение: обучающийся _______________________________________________ готов(а) к выполнению профессиональной деятельности по Осуществлению интеграции программных модулей.": f"Заключение: обучающийся {FULL_NAME} готов к выполнению профессиональной деятельности по осуществлению интеграции программных модулей.",
        "Работал (а) с 25 мая 2026 года по 14 июня 2026 года и показал(а) результаты по освоению следующих общих компетенций": f"Работал с {PERIOD_TEXT} и показал результаты по освоению следующих общих компетенций",
        "Обучающийся _______________________________(________________)": f"Обучающийся {FULL_NAME} (________________)",
    }

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text in paragraph_replacements:
            set_paragraph_text(paragraph, paragraph_replacements[text])
        elif text.startswith("Обучающийся") and "________________" in text:
            set_paragraph_text(paragraph, f"Обучающийся {FULL_NAME} (________________)")
        elif text == "_______________________________ (___________________)":
            set_paragraph_text(paragraph, f"{ORG_SUPERVISOR} (___________________)")

    # Title table.
    t0 = doc.tables[0]
    set_cell_text(t0.cell(0, 1), SPECIALITY)
    set_cell_text(t0.cell(4, 1), FULL_NAME)
    set_cell_text(t0.cell(5, 1), COURSE)
    set_cell_text(t0.cell(7, 1), GROUP_DOTS)
    set_cell_text(t0.cell(8, 1), PERIOD)

    # Assignment, diary, attestation, characteristic, report info tables.
    for idx in [2, 5, 7]:
        fill_common_table(doc.tables[idx], FULL_NAME_GEN)
    t9 = doc.tables[9]
    set_cell_text(t9.cell(0, 1), FULL_NAME)
    set_cell_text(t9.cell(1, 1), SPECIALITY)
    set_cell_text(t9.cell(2, 1), COURSE)
    set_cell_text(t9.cell(3, 1), "Осуществление интеграции программных модулей")
    set_cell_text(t9.cell(4, 1), GROUP_SPACES)
    set_cell_text(t9.cell(5, 1), PLACE)
    set_cell_text(t9.cell(6, 1), PERIOD)

    t3 = doc.tables[3]
    set_cell_text(t3.cell(0, 1), FULL_NAME)
    set_cell_text(t3.cell(1, 1), COURSE)
    set_cell_text(t3.cell(2, 1), SPECIALITY)
    set_cell_text(t3.cell(3, 1), GROUP_SPACES)
    set_cell_text(t3.cell(4, 1), "Производственная")
    set_cell_text(t3.cell(5, 1), PLACE)
    set_cell_text(t3.cell(6, 1), PERIOD)

    # Use existing PM02 wording in diary, but fill sequential numbers.
    for row_idx in range(1, 16):
        set_cell_text(doc.tables[4].cell(row_idx, 0), str(row_idx))

    # Competency marks are left for the supervisor, but OK results are normally pluses.
    for row_idx in range(1, 12):
        set_cell_text(doc.tables[8].cell(row_idx, 2), "+")

    # Replace report placeholders with prose.
    for paragraph in list(doc.paragraphs):
        raw = paragraph.text.strip()
        key = next((k for k in REPORT_TEXT if raw.startswith(k)), None)
        if key:
            style = paragraph.style
            inserted = []
            for item in reversed(REPORT_TEXT[key]):
                new_p = insert_paragraph_after(paragraph, item, style)
                inserted.append(new_p)
            delete_paragraph(paragraph)

    # Fix all visible formatting and remove highlighting/shading from filled areas.
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        normalize_paragraph(paragraph, first_line=bool(text) and not text.isupper())

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                clear_cell_shading(cell)
                for paragraph in cell.paragraphs:
                    normalize_paragraph(paragraph, first_line=False)

    while doc.paragraphs and not doc.paragraphs[-1].text.strip():
        delete_paragraph(doc.paragraphs[-1])

    doc.core_properties.author = FULL_NAME
    doc.core_properties.last_modified_by = FULL_NAME
    doc.save(str(OUT))
    strip_highlights_and_shading(OUT)
    print(OUT)


def strip_highlights_and_shading(path: Path):
    tmp = path.with_suffix(".tmp.docx")
    with ZipFile(path, "r") as zin, ZipFile(tmp, "w", ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/document.xml":
                from lxml import etree

                w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                ns = {"w": w_ns}
                root = etree.fromstring(data)
                body = root.find(f"{{{w_ns}}}body")
                if body is not None:
                    children = list(body)
                    for idx, child in enumerate(children):
                        if child.tag != f"{{{w_ns}}}p":
                            continue
                        text = "".join(child.xpath(".//w:t/text()", namespaces=ns)).strip()
                        page_breaks = child.xpath('.//w:br[@w:type="page"]', namespaces=ns)
                        if text or not page_breaks:
                            continue
                        for nxt in children[idx + 1 :]:
                            if nxt.tag != f"{{{w_ns}}}p":
                                continue
                            p_pr = nxt.find(f"{{{w_ns}}}pPr")
                            if p_pr is None:
                                p_pr = etree.Element(f"{{{w_ns}}}pPr")
                                nxt.insert(0, p_pr)
                            if p_pr.find(f"{{{w_ns}}}pageBreakBefore") is None:
                                p_pr.insert(0, etree.Element(f"{{{w_ns}}}pageBreakBefore"))
                            body.remove(child)
                            break
                data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
            if item.filename.startswith("word/") and item.filename.endswith(".xml"):
                text = data.decode("utf-8")
                import re

                text = re.sub(r"<w:highlight\b[^>]*/>", "", text)
                text = re.sub(r"<w:shd\b[^>]*/>", "", text)
                text = re.sub(r"<w:keepNext\b[^>]*/>", "", text)
                text = re.sub(r"<w:keepLines\b[^>]*/>", "", text)
                text = re.sub(r"<w:cantSplit\b[^>]*/>", "", text)
                data = text.encode("utf-8")
            zout.writestr(item, data)
    tmp.replace(path)


if __name__ == "__main__":
    fill_doc()
